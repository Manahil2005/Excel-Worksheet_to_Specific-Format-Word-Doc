#!/usr/bin/env python3
"""
Generate the weekly "BRIEF ON WEEKLY INFLATION" Word summary directly from
the SPI worksheet (xlsx), matching the format of Summary_23_07_2026.doc.

USAGE (each week):
    python3 generate_spi_summary.py --excel Worksheet_DD_MM_YYYY.xlsx --output Summary_DD_MM_YYYY.docx

It reads:
  - "Impact Combined" sheet  -> combined SPI % change (week-on-week & year-on-year),
                                 group weights/impacts (Food, Non-Food, Utilities, Transport)
  - "Annexure-V" sheet       -> this week's % change by consumption quintile (Q1-Q5, Combined)
  - "Sorted - MoM and YoY"   -> per-item % change lists used for the narrative paragraphs

It also reads/writes "spi_history.csv" (kept alongside this script) which stores the
weekly quintile % change for every week processed. Table 1 of the report ("current week
over previous 5 weeks") needs that history -- a single week's workbook only contains the
current week's number, so the history file is what carries the other 5 columns forward.
Keep spi_history.csv next to the script and reuse it every week; the script appends the
new week to it automatically.

Requires: openpyxl, python-docx  (both already available in this environment)
"""
import argparse
import csv
import os
import re
import sys
from datetime import datetime

import openpyxl
import docx

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "spi_template.docx")
HISTORY_PATH = os.path.join(SCRIPT_DIR, "spi_history.csv")

# "Major increase" cutoff used for the Year-on-Year narrative paragraph only.
# (Observed convention in the original report: all YoY *decreases* are listed,
#  but only YoY *increases* at/above this threshold are called out by name.
#  Week-on-week narrative lists every non-zero item, so no threshold there.)
YOY_INCREASE_THRESHOLD = 13.0

# Short/editorial names used in the narrative sentences, as seen in the source report.
# Items not listed here fall back to their full worksheet description.
SHORT_NAMES = {
    "Wheat Flour Bag": "Wheat Flour",
    "Rice Basmati Broken (Average Quality)": "Rice Basmati",
    "Rice IRRI-6/9 (Sindh/Punjab)": "IRRI",
    "Bread plain (Small Size)": "Bread",
    "Beef with Bone (Average Quality)": "Beef",
    "Mutton (Average Quality)": "Mutton",
    "Chicken Farm Broiler (Live)": "Chicken",
    "Milk fresh (Un-boiled)": "Milk Fresh",
    "Curd (Dahi) Loose": "Curd",
    "Powdered Milk NIDO 390 gm Polybag": "Powdered Milk",
    "Eggs Hen (Farm)": "Eggs",
    "Mustard Oil (Average Quality)": "Mustard Oil",
    "Cooking Oil DALDA or Other Similar Brand (SN), 5 Litre Tin": "Cooking Oil",
    "Vegetable Ghee DALDA/HABIB 2.5 kg Tin": "Vegetable Ghee 2.5 KG",
    "Vegetable Ghee DALDA/HABIB or Other superior Quality 1 kg Pouch": "Vegetable Ghee 1Kg",
    "Bananas (Kela) Local": "Bananas",
    "Pulse Masoor (Washed)": "Masoor",
    "Pulse Moong (Washed)": "Moong",
    "Pulse Mash (Washed)": "Mash",
    "Pulse Gram": "Gram",
    "Potatoes": "Potatoes",
    "Onions": "Onions",
    "Tomatoes": "Tomatoes",
    "Sugar Refined": "Sugar",
    "Gur (Average Quality)": "Gur",
    "Salt Powdered (NATIONAL/SHAN) 800 gm Packet": "Salt Powder",
    "Chilies Powder NATIONAL 200 gm Packet": "Chilies Powder",
    "Garlic (Lehsun)": "Garlic",
    "Tea Lipton Yellow Label 190 gm Packet": "Tea Lipton",
    "Cooked Beef at Average Hotel": "Cooked Beef",
    "Cooked Daal at Average Hotel": "Cooked Daal",
    "Tea Prepared Ordinary": "Tea Prepared",
    "Cigarettes Capstan 20'S Packet": "Cigarettes",
    "Long Cloth 57\" Gul Ahmed/Al Karam": "Long Cloth",
    "Shirting (Average Quality)": "Shirting",
    "Lawn Printed Gul Ahmed/Al Karam": "Lawn Printed",
    "Georgette (Average Quality)": "Georgette",
    "Gents Sandal Bata": "Gents Sandal",
    "Gents Sponge Chappal Bata": "Gents Sponge Chappal",
    "Ladies Sandal Bata": "Ladies Sandal",
    "Electricity Charges for Q1*": "Electricity Charges for Q1",
    "Gas Charges for Q1": "Gas Charges",
    "Firewood Whole": "Firewood",
    "Energy Saver Philips 14 Watt": "Energy Saver",
    "Sufi Washing Soap 250 gm Cake": "Washing Soap",
    "Match Box": "Match Box",
    "Petrol Super": "Petrol",
    "Hi-Speed Diesel": "Diesel",
    "LPG 11.67 kg Cylinder": "LPG",
    "Telephone Call Charges": "Telephone Charges",
    "Toilet Soap LIFEBUOY 115 gm": "Toilet Soap",
}


def short_name(item):
    return SHORT_NAMES.get(item, item)


def fmt_date(d):
    return d.strftime("%d.%m.%Y")


def parse_dates(ws_impact):
    """Row 2 of Impact Combined: 'Current Date = 23/07/2026 ,Previous Date = 16/07/2026 ,Corresponding Date = 24/07/2025'"""
    text = ws_impact.cell(row=2, column=1).value
    m = re.search(
        r"Current Date\s*=\s*([\d/]+)\s*,\s*Previous Date\s*=\s*([\d/]+)\s*,\s*Corresponding Date\s*=\s*([\d/]+)",
        text,
    )
    cur, prev, corr = m.groups()
    fmt = "%d/%m/%Y"
    return (
        datetime.strptime(cur.strip(), fmt),
        datetime.strptime(prev.strip(), fmt),
        datetime.strptime(corr.strip(), fmt),
    )


def get_group_and_total(ws_impact):
    """Find the 'Item group' summary column and pull Food/Non-food/Utility/Transport/Total rows."""
    header_row = None
    header_col = None
    for row in ws_impact.iter_rows(min_row=1, max_row=5):
        for c in row:
            if c.value == "Item group":
                header_row, header_col = c.row, c.column
    if header_row is None:
        raise ValueError("Could not find 'Item group' column in Impact Combined sheet")

    weight_col = header_col + 1
    impact_prev_col = header_col + 2
    impact_corr_col = header_col + 3

    groups = {}
    total = None
    for r in range(header_row + 1, ws_impact.max_row + 1):
        label = ws_impact.cell(row=r, column=header_col).value
        if label is None:
            continue
        weight = ws_impact.cell(row=r, column=weight_col).value
        impact_prev = ws_impact.cell(row=r, column=impact_prev_col).value
        impact_corr = ws_impact.cell(row=r, column=impact_corr_col).value
        if label == "Total":
            total = (weight, impact_prev, impact_corr)
        else:
            groups[label] = (weight, impact_corr)
    return groups, total


def get_quintile_weekly_changes(ws_annex):
    """Returns dict Q1..Q5,COMBINED -> % change vs previous week (the 'General' row of each block)."""
    section_names = {
        "Quintile - I": "Q1",
        "Quintile - II": "Q2",
        "Quintile - III": "Q3",
        "Quintile - IV": "Q4",
        "Quintile - V": "Q5",
        "Quintile - VI (Overall/Combined)": "COMBINED",
    }
    out = {}
    for row in ws_annex.iter_rows(min_row=1, max_row=48):
        val = row[0].value
        if isinstance(val, str) and val.strip() in section_names:
            key = section_names[val.strip()]
            general_row = row[0].row + 6  # General is 6 rows below the section header
            pct_change = ws_annex.cell(row=general_row, column=5).value  # col E
            out[key] = pct_change
    return out


def get_total_item_count(ws_appendix):
    """Counts the 51-item basket from Appendix-A (col A = serial no 1..51, col B = description).
    Used as ground truth for the flat/no-change count, since the 'Sorted' sheet sometimes
    silently drops items whose price didn't move (e.g. Telephone Charges, Toilet Soap)."""
    seen = set()
    for row in ws_appendix.iter_rows(min_row=1, max_row=ws_appendix.max_row):
        no, desc = row[0].value, row[1].value
        if isinstance(no, int) and isinstance(desc, str) and desc.strip():
            seen.add(desc.strip())
    return len(seen)


def get_sorted_changes(ws_sorted):
    """Reads the 'Sorted - MoM and YoY' sheet -> (mom_list, yoy_list), each a list of (item, pct)."""
    mom, yoy = [], []
    for row in ws_sorted.iter_rows(min_row=3, max_row=ws_sorted.max_row):
        item_a, pct_a = row[0].value, row[1].value
        item_d, pct_d = row[3].value, row[4].value
        if isinstance(item_a, str) and isinstance(pct_a, (int, float)):
            mom.append((item_a, float(pct_a)))
        if isinstance(item_d, str) and isinstance(pct_d, (int, float)):
            yoy.append((item_d, float(pct_d)))
    return mom, yoy


def build_change_list_text(items, threshold=0.0):
    """items: list of (name, pct). Groups items with an identical rounded pct together
    (e.g. two items both at 0.09% become 'Curd & Milk Fresh (0.09%)'), sorted by
    magnitude descending, and returns the report-style comma list ending in ' and X'."""
    filtered = [(n, p) for n, p in items if abs(p) >= threshold and p != 0]
    filtered.sort(key=lambda x: abs(x[1]), reverse=True)

    # group consecutive items sharing the same rounded percentage
    groups = []
    for name, pct in filtered:
        rp = round(pct, 2)
        if groups and groups[-1][0] == rp:
            groups[-1][1].append(name)
        else:
            groups.append([rp, [name]])

    parts = []
    for rp, names in groups:
        label = " & ".join(short_name(n) for n in names)
        parts.append(f"{label} ({rp:.2f}%)")

    if not parts:
        return "none"
    if len(parts) == 1:
        return parts[0]
    return ", ".join(parts[:-1]) + " and " + parts[-1]


def load_history(path):
    rows = []
    if os.path.exists(path):
        with open(path, newline="", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                rows.append(row)
    return rows


def save_history(path, rows):
    fieldnames = ["date", "Q1", "Q2", "Q3", "Q4", "Q5", "COMBINED"]
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for r in rows:
            w.writerow(r)


def replace_placeholders(doc, mapping):
    def process_paragraph(p):
        for run in p.runs:
            for key, val in mapping.items():
                token = "{{" + key + "}}"
                if token in run.text:
                    run.text = run.text.replace(token, str(val))

    for p in doc.paragraphs:
        process_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--excel", required=True, help="Path to this week's Worksheet_DD_MM_YYYY.xlsx")
    ap.add_argument("--output", required=True, help="Path to write the Summary .docx")
    ap.add_argument("--template", default=TEMPLATE_PATH)
    ap.add_argument("--history", default=HISTORY_PATH)
    args = ap.parse_args()

    wb = openpyxl.load_workbook(args.excel, data_only=True)
    ws_impact = wb["Impact Combined"]
    ws_annex = wb["Annexure-V"]
    ws_sorted = wb["Sorted - MoM and YoY"]
    ws_appendix = wb["Appendix-A"]

    date_curr, date_prev, date_corr = parse_dates(ws_impact)
    groups, total = get_group_and_total(ws_impact)
    weight_total, combined_mom, combined_yoy = total

    quintiles_now = get_quintile_weekly_changes(ws_annex)
    mom_items, yoy_items = get_sorted_changes(ws_sorted)

    mom_inc = [(n, p) for n, p in mom_items if p > 0]
    mom_dec = [(n, p) for n, p in mom_items if p < 0]
    total_items = get_total_item_count(ws_appendix)
    mom_flat_count = total_items - len(mom_inc) - len(mom_dec)

    yoy_inc = [(n, p) for n, p in yoy_items if p > 0]
    yoy_dec = [(n, p) for n, p in yoy_items if p < 0]

    mom_increase_list = build_change_list_text(mom_inc, threshold=0.0)
    mom_decrease_list = build_change_list_text(mom_dec, threshold=0.0)
    yoy_increase_list = build_change_list_text(yoy_inc, threshold=YOY_INCREASE_THRESHOLD)
    yoy_decrease_list = build_change_list_text(yoy_dec, threshold=0.0)

    # ---- history (Table 1: current week vs previous 5 weeks) ----
    history = load_history(args.history)
    history = [r for r in history if r["date"] != fmt_date(date_curr)]  # avoid dup if re-run
    new_row = {"date": fmt_date(date_curr)}
    for q in ["Q1", "Q2", "Q3", "Q4", "Q5", "COMBINED"]:
        new_row[q] = f"{quintiles_now.get(q, 0):.2f}"
    history.append(new_row)
    history.sort(key=lambda r: datetime.strptime(r["date"], "%d.%m.%Y"))
    save_history(args.history, history)

    last6 = history[-6:]
    last6 = list(reversed(last6))  # most recent first, matches report's left-to-right order
    while len(last6) < 6:
        last6.append({"date": "", "Q1": "", "Q2": "", "Q3": "", "Q4": "", "Q5": "", "COMBINED": ""})

    mapping = {
        "DATE_CURR": fmt_date(date_curr),
        "DATE_CORR": date_corr.strftime("%B %-d, %Y") if os.name != "nt" else date_corr.strftime("%B %#d, %Y"),
        "COMBINED_MOM": f"{combined_mom:.2f}",
        "COMBINED_YOY": f"{combined_yoy:.2f}",
        "HISTORY_START_DATE": "February 19, 2026",
        "MOM_INCREASE_LIST": mom_increase_list,
        "MOM_DECREASE_LIST": mom_decrease_list,
        "MOM_INC_COUNT": f"{len(mom_inc):02d}",
        "MOM_DEC_COUNT": f"{len(mom_dec):02d}",
        "MOM_FLAT_COUNT": f"{mom_flat_count:02d}",
        "YOY_INCREASE_LIST": yoy_increase_list,
        "YOY_DECREASE_LIST": yoy_decrease_list,
        "FOOD_WEIGHT": f"{groups['Food'][0]:.2f}",
        "FOOD_IMPACT": f"{groups['Food'][1]:.2f}",
        "NONFOOD_WEIGHT": f"{groups['Non-food'][0]:.2f}",
        "NONFOOD_IMPACT": f"{groups['Non-food'][1]:.2f}",
        "UTIL_WEIGHT": f"{groups['Utility'][0]:.2f}",
        "UTIL_IMPACT": f"{groups['Utility'][1]:.2f}",
        "TRANSPORT_WEIGHT": f"{groups['Transport'][0]:.2f}",
        "TRANSPORT_IMPACT": f"{groups['Transport'][1]:.2f}",
        "TOTAL_WEIGHT": f"{weight_total:.2f}",
        "TOTAL_IMPACT": f"{combined_yoy:.2f}",
    }
    for i, row in enumerate(last6):
        mapping[f"DATE_W{i}"] = row["date"]
        for q in ["Q1", "Q2", "Q3", "Q4", "Q5", "COMBINED"]:
            mapping[f"{q}_W{i}"] = row[q]

    doc = docx.Document(args.template)
    replace_placeholders(doc, mapping)
    doc.save(args.output)
    print(f"Wrote {args.output}")
    print(f"History file: {args.history} ({len(history)} weeks on record)")


if __name__ == "__main__":
    main()
